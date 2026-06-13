from docx import Document
import os

def create_template(name, text, placeholders=None):
    doc = Document()
    doc.add_heading(f'Bloque: {name}', level=1)
    doc.add_paragraph(text)
    if placeholders:
        p = doc.add_paragraph("Detalles: ")
        for ph in placeholders:
            p.add_run(f"\n{ph}: {{{{{ph}}}}}")
    
    path = os.path.join("templates", name)
    doc.save(path)
    print(f"Created {path}")

def main():
    if not os.path.exists("templates"):
        os.makedirs("templates")

    create_template("header.docx", "INFORME TÉCNICO DE LEVANTAMIENTO", ["instance_id"])
    create_template("camara.docx", "Se detectó la instalación de una cámara de videovigilancia.", 
                    ["modelo_equipo", "tipo_equipo", "foto_nodo", "spec_marca", "spec_descripcion", "spec_caracteristicas"])
    create_template("ups.docx", "Sistema de Respaldo de Energía (UPS) encontrado.", ["modelo_equipo"])
    create_template("cierre.docx", "Fin del reporte generado automáticamente.")

if __name__ == "__main__":
    main()
