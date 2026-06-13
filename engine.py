try:
    import yaml
except ImportError:
    yaml = None
import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Relative import for Odoo package structure
from utils import download_and_resize_image, get_address_from_coords, get_static_map_image

try:
    from docx.enum.table import WD_ALIGN_VERTICAL
except ImportError:
    pass

class ReportEngine:
    def __init__(self, mapping_config=None, catalog_data=None, template_dir="templates", auth_token=None):
        """
        Modified init to accept dicts directly from Odoo models.
        """
        self.template_dir = template_dir
        self.auth_token = auth_token
        
        # Load Mapping
        if isinstance(mapping_config, dict):
            self.mapping_config = mapping_config
        elif isinstance(mapping_config, str):
            # Check for JSON first
            json_path = mapping_config.replace('.yaml', '.json').replace('.yml', '.json')
            if os.path.exists(json_path):
                 with open(json_path, 'r', encoding='utf-8') as f:
                    self.mapping_config = json.load(f)
            elif os.path.exists(mapping_config):
                if yaml:
                    with open(mapping_config, 'r', encoding='utf-8') as f:
                        self.mapping_config = yaml.safe_load(f)
                else:
                    self.mapping_config = {}
                    print("Error: PyYAML missing and no JSON mapping found.")
            else:
                self.mapping_config = {}
        else:
            self.mapping_config = {}
        
        # Load Catalog
        if isinstance(catalog_data, dict):
            self.catalog = catalog_data
        elif isinstance(catalog_data, str) and os.path.exists(catalog_data):
            with open(catalog_data, 'r', encoding='utf-8') as f:
                self.catalog = json.load(f)
        else:
            self.catalog = {}

    def generate_report(self, row_data, output_path="reporte_generado.docx", report_type="memoria_tecnica", schema=None, choices_dict=None):
        """
        Generates a DOCX report based on a single row of data (dict).
        Param `report_type`: Key in mapping.yaml -> report_types
        """
        self.choices_dict = choices_dict or {}
        # Store attachments context for image resolution
        self.current_attachments = row_data.get('_attachments', [])
        
        # 0. Enrich Data with Catalog Specs and Geo
        self._inject_standard_mapping(row_data) # NEW: Map generic keys like 'cliente'
        self._inject_catalog_specs(row_data)
        self._resolve_geo_data(row_data)

        # Get Configuration for this Report Type
        report_config = self.mapping_config.get('report_types', {}).get(report_type, None)
        
        if not report_config:
            print(f"Warning: Configuration for {report_type} not found. Using defaults.")
            template_base = "header.docx"
            blocks_config = self.mapping_config.get('blocks', []) # Fallback to root blocks if any
        else:
            template_base = report_config.get('template_base', "header.docx")
            blocks_config = report_config.get('blocks', [])

        # 1. Create Base Document
        # In Odoo, template_dir should be absolute path
        header_path = os.path.join(self.template_dir, template_base)
        if os.path.exists(header_path):
            doc = Document(header_path)
            # Initial replacement in header
            self._replace_text_in_doc(doc, row_data)
            self._replace_tables_in_doc(doc, row_data)
            self._replace_technical_details(doc, row_data, schema)
        else:
            print(f"Warning: Base template not found at {header_path}")
            doc = Document() # Fallback empty
        
        # 2. Logic to detect blocks (Dynamic based on Report Type)
        if blocks_config:
            for block_rule in blocks_config:
                variable = block_rule['variable'] # e.g. "tipo_equipo"
                expected_value = block_rule['value'] # e.g. "Cámara"
                template_name = block_rule['template'] # e.g. "camara.docx"
                
                actual_value = row_data.get(variable, "")
                
                # Check condition (Case insensitive)
                if str(expected_value).lower() in str(actual_value).lower():
                    self._append_template(doc, template_name, row_data)

        # 4. Save
        doc.save(output_path)
        return output_path

    def _inject_standard_mapping(self, row_data):
        config = self.mapping_config.get('config', {})
        def get_val(path): return row_data.get(path, "")

        if 'column_cliente' in config: row_data['cliente'] = get_val(config['column_cliente'])
        if 'column_proyecto' in config: row_data['proyecto_id'] = get_val(config['column_proyecto'])
        if 'column_sitio' in config: row_data['nombre_sitio'] = get_val(config['column_sitio'])
            
        if not row_data.get('descripcion_breve'):
            row_data['descripcion_breve'] = "Solución Integral de Seguridad y Monitoreo"

    def _inject_catalog_specs(self, row_data):
        model_key = row_data.get("modelo_equipo", "")
        item = self.catalog.get(model_key)
        
        if not item:
            row_data["spec_marca"] = "N/A"
            row_data["spec_descripcion"] = "Especificación no encontrada en catálogo"
            row_data["spec_caracteristicas"] = ""
            return

        row_data["spec_marca"] = item.get("marca", "")
        row_data["spec_descripcion"] = item.get("descripcion", "")
        features = item.get("caracteristicas", [])
        if isinstance(features, list):
            formatted_features = "\n".join([f"• {f}" for f in features])
            row_data["spec_caracteristicas"] = formatted_features
        else:
            row_data["spec_caracteristicas"] = str(features)

    def _resolve_geo_data(self, row_data):
        geo_keys = [k for k, v in row_data.items() if isinstance(v, str) and len(str(v).split(' ')) >= 2]
        if geo_keys:
            val = row_data[geo_keys[0]]
            try:
                parts = val.split(' ')
                lat, lon = parts[0], parts[1]
                address = get_address_from_coords(lat, lon)
                if address: row_data['direccion'] = address
                
                map_bytes = get_static_map_image(lat, lon)
                if map_bytes: row_data['mapa_estatico'] = map_bytes
            except: pass

    def _replace_technical_details(self, doc, data, schema):
        if not schema: return
        target_p = None
        for p in doc.paragraphs:
             if "{{DETALLE_TECNICO}}" in p.text:
                 target_p = p
                 break
        
        if target_p:
            # Table with 3 columns (Grid Layout)
            table = doc.add_table(rows=0, cols=3)
            # Apply Double Borders
            self._set_table_borders(table)
            # Page width ~6.5 inches. 3 cols => ~2.16 inches each.
            for col in table.columns: col.width = Inches(2.15)
                
            self._render_schema_to_table(table, schema.get('children', []), data)
            
            parent = target_p._element.getparent()
            parent.insert(parent.index(target_p._element) + 1, table._element)
            parent.remove(target_p._element)

    def _resolve_image_helper(self, val_str):
        if not val_str: return None, None
        img_url = val_str
        auth_header = None
        
        if not val_str.startswith("http") and hasattr(self, 'current_attachments'):
            found_att = None
            for att in self.current_attachments:
                if att.get('filename') == val_str or att.get('id') == val_str:
                    found_att = att
                    break
            if not found_att:
                suffix = f"/{val_str}"
                for att in self.current_attachments:
                    if att.get('filename', '').endswith(suffix):
                        found_att = att
                        break
            if found_att:
                # Optimization: Prefer reduced size images for faster report generation
                raw_url = found_att.get('medium_download_url') or found_att.get('small_download_url') or found_att.get('download_url')
                if raw_url:
                    img_url = f"https://api.ona.io{raw_url}" if raw_url.startswith("/") else raw_url
                    auth_header = self.auth_token

        if "ona.io" in str(img_url) and self.auth_token:
            auth_header = self.auth_token
        return img_url, auth_header

    def _set_cell_bg(self, cell, color_hex):
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
        cell._element.get_or_add_tcPr().append(shading_elm)

        cell._element.get_or_add_tcPr().append(shading_elm)

    def _set_table_borders(self, table):
        from docx.oxml.shared import OxmlElement, qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Remove existing borders if any
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
            
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'double')
            border.set(qn('w:sz'), '4') 
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        tblPr.append(borders)

    def _set_cell_style(self, cell, text, font_name="Century Gothic", font_size_pt=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, color_hex=None):
        try:
            from docx.enum.table import WD_ALIGN_VERTICAL
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except: pass
        
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        pf = paragraph.paragraph_format
        pf.left_indent = Pt(0)
        pf.right_indent = Pt(0) # Fix right indent
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(2)
        pf.space_after = Pt(2)
        
        run = paragraph.add_run(str(text))
        font = run.font
        font.name = font_name
        font.size = Pt(font_size_pt)
        font.bold = bold
        font.italic = italic
        
        if color_hex:
            from docx.shared import RGBColor
            h = color_hex.lstrip('#')
            rgb = tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
            font.color.rgb = RGBColor(*rgb)

    def _has_visible_data(self, children, data):
        """Recursively check if a group/section has any non-empty fields to display."""
        if not children or not data: return False
        
        for child in children:
            c_type = child.get('type')
            c_name = child.get('name')
            
            if c_type == 'group':
                # Check nested group
                if self._has_visible_data(child.get('children', []), data):
                    return True
            elif c_type == 'repeat':
                # Check repeated items
                repeats = data.get(c_name, [])
                if isinstance(repeats, list) and repeats:
                    # Check if ANY repeat item has visible data
                    for item in repeats:
                        if self._has_visible_data(child.get('children', []), item):
                            return True
            else:
                # Check simple field
                val = self._find_val_fuzzy(data, c_name)
                if val is not None and str(val).strip() != "" and str(val).lower() != 'nan':
                    return True
        return False

    def _render_schema_to_table(self, table, children, data, level=0):
        # Flatten the current level items to render them in a grid
        current_row_items = []
        
        def flush_row_items():
            if not current_row_items: return
            # Add a row for these items
            row = table.add_row()
            for idx, item in enumerate(current_row_items):
                if idx >= 3: break 
                cell = row.cells[idx]
                _render_cell_content(cell, item)
            current_row_items.clear()

        def _render_cell_content(cell, child):
            c_name = child.get('name')
            c_label = child.get('label', c_name)
            c_type = child.get('type', 'text')
            
            # 1. Render Label (Bold Header)
            self._set_cell_bg(cell, "F8F9FA")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"{c_label}")
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(44, 62, 80) 
            
            # 2. Render Value (Body)
            val = self._find_val_fuzzy(data, c_name)
            if val is None or str(val).strip() == "" or str(val).lower() == 'nan':
                p.add_run("\nN/A")
                return

            # Check for Geopoint Heuristic First (Strict)
            val_str = str(val)
            parts = val_str.split()
            is_geo = False
            # Must have 2-4 parts, digits, AND contain a decimal point to avoid phone numbers
            if len(parts) >= 2 and len(parts) <= 4 and any('.' in p for p in parts):
                try:
                    lat = float(parts[0])
                    lon = float(parts[1])
                    if -90 <= lat <= 90 and -180 <= lon <= 180 and (abs(lat) > 0.001 or abs(lon) > 0.001):
                        is_geo = True
                        p.add_run("\n")
                        map_stream = get_static_map_image(lat, lon, zoom=15, width=400, height=300)
                        if map_stream:
                            run = p.add_run()
                            run.add_picture(map_stream, width=Inches(1.8))
                        p.add_run(f"\n{val_str}")
                except ValueError: pass
            
            if is_geo: return

            # Formatting by Type
            if c_type in ['photo', 'image']:
                p.add_run("\n")
                img_url, auth_header = self._resolve_image_helper(val_str)
                # Resize to ~1.8 inches 
                img_stream = download_and_resize_image(img_url, target_width=300, auth_token=auth_header)
                if img_stream:
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(1.8))
                else:
                    p.add_run(val_str)
            
            # Select / Labels
            elif c_type and (c_type.startswith('select_one') or c_type.startswith('select_multiple')):
                # Resolve Labels logic
                list_name = child.get('list_name') or child.get('itemset') or child.get('select_from_list_name')
                if not list_name:
                    parts = c_type.split()
                    if len(parts) > 1: list_name = parts[-1]

                val_parts = val_str.split(' ')
                resolved = []
                for part in val_parts:
                    found = part
                    # Try global dict first
                    if self.choices_dict and list_name and list_name in self.choices_dict:
                        found = self.choices_dict[list_name].get(part, part)
                    # Try local embedded children
                    elif 'children' in child:
                        for ch in child.get('children', []):
                            if ch.get('name') == part:
                                found = ch.get('label', part); break
                    resolved.append(str(found))
                p.add_run(f"\n{', '.join(resolved)}")

            else:
                # Fallback: Check if it's a select field disguised as text (Has list_name?)
                list_name = child.get('list_name') or child.get('itemset') or child.get('select_from_list_name')
                if list_name and self.choices_dict:
                     # Try to resolve label
                     val_str = str(val)
                     if list_name in self.choices_dict:
                         resolved_val = self.choices_dict[list_name].get(val_str, val_str)
                         p.add_run(f"\n{resolved_val}")
                     else:
                         p.add_run(f"\n{val_str}")
                else:
                    p.add_run(f"\n{val_str}")


        # Main Loop for Children
        for child in children:
            c_type = child.get('type')
            
            # If Container (Group/Repeat), flush current row and take full width
            if c_type in ['group', 'repeat']:
                # Optimization: Check if this group/repeat has any data TO SHOW. 
                # If not, skip the header entirely.
                has_data = False
                if c_type == 'group':
                    has_data = self._has_visible_data(child.get('children', []), data)
                elif c_type == 'repeat':
                    repeats = data.get(child.get('name'), [])
                    if isinstance(repeats, list) and repeats:
                         for item in repeats:
                             if self._has_visible_data(child.get('children', []), item):
                                 has_data = True; break
                
                if not has_data:
                    continue

                flush_row_items()
                
                # Render Section Header - Merge all 3 columns
                c_label = child.get('label', child.get('name'))
                row = table.add_row()
                cell = row.cells[0]
                cell.merge(row.cells[1])
                cell.merge(row.cells[2])
                self._set_cell_style(cell, c_label, bold=True, color_hex="#16a085", font_size_pt=10)
                self._set_cell_bg(cell, "E8F6F3")
                
                if c_type == 'group':
                    self._render_schema_to_table(table, child.get('children', []), data, level+1)
                elif c_type == 'repeat':
                    repeats = data.get(child.get('name'), [])
                    if isinstance(repeats, list):
                        for i, item in enumerate(repeats):
                            # Item Header - Merge all 3 columns
                            r_row = table.add_row()
                            r_cell = r_row.cells[0]
                            r_cell.merge(r_row.cells[1])
                            r_cell.merge(r_row.cells[2])
                            self._set_cell_style(r_cell, f"   Item #{i+1}", italic=True, font_size_pt=9)
                            self._render_schema_to_table(table, child.get('children', []), item, level+1)
            else:
                # Accumulate simple fields
                # Filter out empty/skipped fields
                c_name = child.get('name')
                val = self._find_val_fuzzy(data, c_name)
                if val is None or str(val).strip() == "" or str(val).lower() == 'nan':
                     continue
                
                current_row_items.append(child)
                if len(current_row_items) == 3:
                    flush_row_items()
        
        # Flush remaining
        flush_row_items()

    def _append_template(self, main_doc, template_name, data):
        template_path = os.path.join(self.template_dir, template_name)
        if not os.path.exists(template_path):
            print(f"Warning: Template {template_name} not found.")
            return

        sub_doc = Document(template_path)
        self._replace_text_in_doc(sub_doc, data)
        self._replace_images_in_doc(sub_doc, data)
        self._replace_tables_in_doc(sub_doc, data)

        for element in sub_doc.element.body:
            main_doc.element.body.append(element)

    def _replace_text_in_doc(self, doc, data):
        for p in doc.paragraphs:
            if params := self._find_placeholders(p.text):
                for key in params:
                    if key in data and not self._is_image_field(key, data):
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in p.text:
                            p.text = p.text.replace(placeholder, str(data[key]))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if params := self._find_placeholders(p.text):
                            for key in params:
                                if key in data and not self._is_image_field(key, data):
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in p.text:
                                        p.text = p.text.replace(placeholder, str(data[key]))

    def _replace_images_in_doc(self, doc, data):
        def process_paragraph_for_images(paragraph):
            if params := self._find_placeholders(paragraph.text):
                for key in params:
                    if self._is_image_field(key, data):
                        val_str = str(data[key])
                        placeholder = f"{{{{{key}}}}}"
                        
                        image_files = [v.strip() for v in val_str.split(',')] if ',' in val_str else [val_str]
                        
                        success_images = 0
                        for img_val in image_files:
                            if not img_val: continue
                            img_url, auth_header = self._resolve_image_helper(img_val)
                            if img_url and str(img_url).startswith("http"):
                                image_stream = download_and_resize_image(img_url, auth_token=auth_header)
                                if image_stream:
                                    if success_images == 0:
                                        paragraph.text = paragraph.text.replace(placeholder, "")
                                    run = paragraph.add_run()
                                    run.add_picture(image_stream, width=Inches(3.5))
                                    run.add_text("  ") # Space between multiple images
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    success_images += 1
                                    
                        if success_images == 0 and placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, "")
                            
        for p in doc.paragraphs: process_paragraph_for_images(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: process_paragraph_for_images(p)

    def _replace_tables_in_doc(self, doc, data):
        table_data = data.get('tabla_cotizacion', [])
        if not table_data:
            return

        target_p = None
        for p in doc.paragraphs:
            if "{{TABLE_PROPOSAL}}" in p.text:
                target_p = p
                break
        if not target_p:
            for outer_table in doc.tables:
                for row in outer_table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if "{{TABLE_PROPOSAL}}" in p.text:
                                target_p = p
                                break
                        if target_p: break
                    if target_p: break
                if target_p: break
        
        if target_p:
            target_p.text = target_p.text.replace("{{TABLE_PROPOSAL}}", "")
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            widths = [Inches(0.7), Inches(3.8), Inches(1.2), Inches(1.2)]
            
            hdr_cells = table.rows[0].cells
            headers = ["CANT", "DESCRIPCIÓN", "C.U MXN", "C.T MXN"]
            purple_hex = "7030A0"
            for j, h in enumerate(headers):
                cell = hdr_cells[j]
                cell.width = widths[j]
                self._set_cell_bg(cell, purple_hex)
                self._set_cell_style(cell, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color_hex="FFFFFF")
            
            grand_subtotal = 0.0
            for i, item in enumerate(table_data):
                row_cells = table.add_row().cells
                try: qty = float(item.get("Cant.", 0))
                except: qty = 0
                cost_str = str(item.get("Costo Unitario", "0")).replace("$", "").replace(",", "")
                try: cost = float(cost_str)
                except: cost = 0.0
                subtotal = qty * cost
                grand_subtotal += subtotal
                
                c = row_cells[0]; c.width = widths[0]
                fmt_qty = int(qty) if qty.is_integer() else qty
                self._set_cell_style(c, fmt_qty, align=WD_ALIGN_PARAGRAPH.CENTER)
                c = row_cells[1]; c.width = widths[1]
                clean_desc = str(item.get("Descripción", "")).replace("**", "")
                self._set_cell_style(c, clean_desc, align=WD_ALIGN_PARAGRAPH.LEFT)
                c = row_cells[2]; c.width = widths[2]
                self._set_cell_style(c, f"${cost:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT)
                c = row_cells[3]; c.width = widths[3]
                self._set_cell_style(c, f"${subtotal:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT)
            
            # Footer rows (Subtotal, IVA, Total)
            row_sub = table.add_row()
            row_sub.cells[0].merge(row_sub.cells[1])
            c_lbl = row_sub.cells[2]; c_lbl.width = widths[2]
            self._set_cell_bg(c_lbl, purple_hex)
            self._set_cell_style(c_lbl, "Subtotal", bold=True, color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
            c_val = row_sub.cells[3]; c_val.width = widths[3]
            self._set_cell_style(c_val, f"${grand_subtotal:,.2f}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
            
            iva_rate = 0.16; iva_val = grand_subtotal * iva_rate
            row_iva = table.add_row(); row_iva.cells[0].merge(row_iva.cells[1])
            c_lbl = row_iva.cells[2]; self._set_cell_bg(c_lbl, purple_hex)
            self._set_cell_style(c_lbl, "IVA 16%", bold=True, color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
            c_val = row_iva.cells[3]; self._set_cell_style(c_val, f"${iva_val:,.2f}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
            
            total_val = grand_subtotal + iva_val
            row_tot = table.add_row(); row_tot.cells[0].merge(row_tot.cells[1])
            c_lbl = row_tot.cells[2]; self._set_cell_bg(c_lbl, purple_hex)
            self._set_cell_style(c_lbl, "TOTAL", bold=True, color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.LEFT)
            c_val = row_tot.cells[3]; self._set_cell_style(c_val, f"${total_val:,.2f}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

            parent = target_p._element.getparent()
            parent.insert(parent.index(target_p._element) + 1, table._element)
            parent.remove(target_p._element)

    def _find_placeholders(self, text):
        import re
        return re.findall(r'\{\{(.*?)\}\}', text)
        
    def _find_val_fuzzy(self, data, key):
        if key in data: return data[key]
        suffix = f"/{key}"
        for k, v in data.items():
            if k.endswith(suffix): return v
        return None

    def _is_image_field(self, key, data):
        val = str(data.get(key, ""))
        if val.startswith("http") and any(ext in val.lower() for ext in ['.jpg', '.jpeg', '.png']): return True
        if "foto" in key.lower() or "image" in key.lower(): return True
        return False
